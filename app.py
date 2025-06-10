import streamlit as st
from openai import OpenAI
import os
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Asistente Pedagógico IA", layout="centered")
st.title("🎓 Asistente Pedagógico con IA")

oai_key = st.text_input("🔑 Ingresa tu API Key de OpenAI:", type="password")

tab1, tab2 = st.tabs(["Generar Sesión", "Acerca del Proyecto"])

with tab1:
    with st.form("formulario_sesion"):
        nivel = st.selectbox("📚 Nivel educativo:", ["Primaria", "Secundaria", "Técnico"])
        competencia = st.text_input("🎯 Competencia o estándar de aprendizaje:")
        tema = st.text_input("🧠 Tema o contenido específico:")
        generar = st.form_submit_button("✨ Generar Sesión")

    if generar:
        if not oai_key:
            st.error("Por favor, ingresa tu API Key de OpenAI.")
        else:
            with st.spinner("Generando sesión con IA..."):
                prompt = f"""
Eres un experto pedagógico. Crea una sesión de aprendizaje para el nivel {nivel} basada en la competencia: {competencia}, y el tema: {tema}.
Incluye título, propósito, actividades de inicio, desarrollo y cierre, evaluación y recursos.
"""
                try:
                    client = OpenAI(api_key=oai_key)
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "Eres un asistente pedagógico altamente calificado."},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    sesion = response.choices[0].message.content
                    st.success("✅ Sesión generada con éxito")
                    st.text_area("📄 Sesión generada:", sesion, height=400)

                    doc = Document()
                    doc.add_heading("Sesión de Aprendizaje IA", 0)
                    doc.add_paragraph(sesion)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        label="📥 Descargar sesión en Word",
                        data=buffer,
                        file_name="sesion_IA.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"❌ Error al generar sesión: {e}")

with tab2:
    st.markdown("""
    ### 🤖 ¿Qué hace este asistente?
    Esta app te permite generar sesiones pedagógicas personalizadas usando IA. Solo ingresa el nivel, competencia y tema. 

    - Desarrollado con Streamlit y OpenAI
    - Ideal para docentes innovadores
    - MVP de la plataforma "Mi Asistente Docente IA"
    """)
